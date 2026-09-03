import json
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "1F4E79"
DARK_BLUE = "17365D"
GREEN = "70AD47"
LIGHT_GREEN = "E2F0D9"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
RED = "9C0006"
GOLD = "806000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

TABLE_PURPOSES = {
    "ai_predictions": "Stores structured outputs from Gemini and future fallback/shadow models for each inspection.",
    "ai_usage": "Tracks provider/model requests, latency, token/image usage, success and estimated cost.",
    "app_settings": "Stores centrally managed application policies and feature/model settings.",
    "approval_requests": "Queues proposed governed changes and their approval decisions.",
    "audit_logs": "Keeps immutable-style before/after records of important administrative actions.",
    "crops": "Controlled crop names and catalogue aliases used for deterministic matching.",
    "employee_roles": "Many-to-many assignment of system roles to employees, including grant audit data.",
    "employees": "Employee identity, contact, Microsoft sign-in linkage, reporting hierarchy and access status.",
    "expert_reviews": "Stores agronomist verification/correction of AI results and dataset eligibility.",
    "inspection_images": "Stores S3 object paths and image metadata; image binary data is not stored in PostgreSQL.",
    "inspection_recommendations": "Records which approved CLSL products were displayed, in what order and why.",
    "inspections": "The central field-inspection record linking employee, crop, location, symptoms and processing state.",
    "organizational_units": "Models company, zone, state, region and territory hierarchy with manager ownership.",
    "problems": "Controlled disease, pest, weed, deficiency and stress vocabulary with aliases.",
    "product_categories": "Controlled CLSL product categories.",
    "product_crop_mappings": "Approved catalogue-explicit product-to-crop relationships.",
    "product_problem_mappings": "Approved catalogue-explicit product-to-problem relationships.",
    "products": "Official CLSL catalogue product master and approval metadata.",
    "roles": "Role definitions used by role-based access control.",
    "schema_migrations": "Applied database migration versions and checksums.",
    "weather_snapshots": "Weather evidence and farm-activity guidance captured with an inspection.",
}

COLUMN_DESCRIPTIONS = {
    "id": "Stable UUID identifier for this record.",
    "employee_code": "Official CLSL employee code; unique.",
    "full_name": "Employee's full name from the approved employee source.",
    "office_email": "Company email used to match Microsoft sign-in.",
    "microsoft_upn": "Microsoft Entra user principal name expected at sign-in.",
    "microsoft_subject": "Immutable Microsoft identity subject after first verified login.",
    "microsoft_tenant_id": "Microsoft Entra tenant identifier after verified login.",
    "reporting_manager_id": "Resolved employee record for the reporting manager.",
    "reporting_manager_name": "Source manager name retained for import traceability.",
    "organizational_unit_id": "Assigned company/zone/state/region/territory record.",
    "source_row": "Original employee workbook row number for traceability.",
    "preferred_language": "Language code selected for the application interface and results.",
    "status": "Controlled lifecycle/access state for the record.",
    "approval_status": "Pending, approved or rejected governance state.",
    "approved_by": "Employee who approved the governed record.",
    "approved_at": "Time the record was approved.",
    "submitted_by": "Employee who proposed the mapping.",
    "aliases": "JSON list of alternate catalogue spellings/names.",
    "catalogue_version": "Official catalogue edition used as the source.",
    "source_page": "Page number in the official catalogue.",
    "storage_key": "Private S3 object key; no image binary is stored in this table.",
    "retention_status": "Whether the image is retained, archived or deleted.",
    "consent_for_training": "Whether this image may be included in a training dataset.",
    "sha256": "Optional content checksum for duplicate/integrity checking.",
    "prediction_role": "Whether the model output is primary, fallback or silent shadow output.",
    "raw_response": "Structured JSON returned by the external/internal service for auditability.",
    "dataset_eligible": "Expert decision that the reviewed case can enter a curated dataset.",
    "match_reason": "Catalogue-grounded explanation for showing the product.",
    "match_score": "Deterministic ranking score used by the product engine.",
    "estimated_cost_usd": "Estimated provider cost for operational reporting.",
    "proposed_data": "JSON representation of the proposed governed change.",
    "previous_data": "JSON snapshot before an audited change.",
    "new_data": "JSON snapshot after an audited change.",
    "value": "JSON setting value.",
    "checksum": "SHA-256 checksum used to detect changed applied migrations.",
}


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, data_rows, widths_dxa, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    set_repeat_header(table.rows[0])
    for row_index, values in enumerate(data_rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2:
                shade(cells[idx], "F8FAFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run("—" if value in (None, "") else str(value))
            set_run_font(r, size=font_size, color="202124")
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_status_callout(doc, title, text, positive=True):
    fill = LIGHT_GREEN if positive else "FFF2CC"
    border_color = GREEN if positive else "BF9000"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color="375623" if positive else GOLD)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Inches(0.12)
    p2_pr = p2._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:fill"), fill)
    p2_pr.append(shd2)
    borders2 = OxmlElement("w:pBdr")
    left2 = OxmlElement("w:left")
    left2.set(qn("w:val"), "single")
    left2.set(qn("w:sz"), "18")
    left2.set(qn("w:space"), "8")
    left2.set(qn("w:color"), border_color)
    borders2.append(left2)
    p2_pr.append(borders2)
    set_run_font(p2.add_run(text), size=10)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def type_label(column):
    data_type = column["data_type"]
    if data_type == "USER-DEFINED":
        return column["udt_name"]
    if data_type == "character varying" and column.get("character_maximum_length"):
        return f"varchar({column['character_maximum_length']})"
    if data_type == "numeric" and column.get("numeric_precision"):
        return f"numeric({column['numeric_precision']},{column['numeric_scale']})"
    return data_type.replace("timestamp with time zone", "timestamptz")


def describe_column(column):
    name = column["column_name"]
    if name in COLUMN_DESCRIPTIONS:
        return COLUMN_DESCRIPTIONS[name]
    if name.endswith("_id"):
        return f"Reference to the related {name[:-3].replace('_', ' ')} record."
    if name.endswith("_at"):
        return f"Timestamp for {name[:-3].replace('_', ' ')}."
    if name.endswith("_count"):
        return f"Recorded number of {name[:-6].replace('_', ' ')}."
    return name.replace("_", " ").capitalize() + "."


def main(audit_path: str, output_path: str, logo_path: str) -> None:
    audit = json.loads(Path(audit_path).read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "CROP LIFE AI  |  DATABASE DICTIONARY"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Editorial-cover header pattern, adapted for a technical reference guide.
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    if Path(logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(logo_path, width=Inches(1.75))
        picture._inline.docPr.set("descr", "Crop Life Science Limited logo")
        picture._inline.docPr.set("title", "CLSL logo")
        p.paragraph_format.space_after = Pt(24)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("PRODUCTION DATA REFERENCE"), size=10, bold=True, color=GREEN)
    kicker.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("Crop Life AI"), size=30, bold=True, color=DARK_BLUE)
    title.paragraph_format.space_after = Pt(3)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("PostgreSQL Database Dictionary & Data Audit"), size=16, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(20)
    add_table(
        doc,
        ["Database", "Snapshot", "Catalogue", "Employees"],
        [[audit["database"]["name"], "3 September 2026", "CLSL English 2023", audit["employee_summary"]["total"]]],
        [2340, 2340, 2340, 2340],
        font_size=9.5,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    set_run_font(p.add_run("Crop Life Science Limited"), size=12, bold=True, color=DARK_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p2.add_run("Prepared for controlled internal storage and implementation reference"), size=9.5, italic=True, color=MID_GRAY)
    doc.add_page_break()

    add_heading(doc, "1. Audit outcome", 1)
    integrity = audit["integrity"]
    integrity_ok = all(int(value) == 0 for value in integrity.values())
    add_status_callout(
        doc,
        "AUDIT PASSED" if integrity_ok else "AUDIT REQUIRES ATTENTION",
        "All automated integrity, approval-ownership, duplicate, email-domain and role-assignment checks passed."
        if integrity_ok
        else "One or more automated integrity checks did not pass; see the integrity table below.",
        positive=integrity_ok,
    )
    add_body(doc, "The database was checked against the employee import, the 73-product CLSL catalogue data and the production modules defined in the master project report. Missing operational modules were added through migration 002. One employee self-manager source entry was normalized to no manager so the hierarchy cannot contain an invalid self-reference.")
    summary_rows = [
        ("Employees", audit["employee_summary"]["total"]),
        ("Active Microsoft-login records", audit["employee_summary"]["active"]),
        ("Pending office email", audit["employee_summary"]["pending"]),
        ("Products", audit["counts"]["products"]),
        ("Product categories", audit["counts"]["product_categories"]),
        ("Crops", audit["counts"]["crops"]),
        ("Problems", audit["counts"]["problems"]),
        ("Approved crop mappings", audit["counts"]["product_crop_mappings"]),
        ("Approved problem mappings", audit["counts"]["product_problem_mappings"]),
        ("Database tables", len(audit["tables"])),
    ]
    add_table(doc, ["Verified item", "Count"], summary_rows, [7020, 2340], font_size=9.5)

    add_heading(doc, "2. Connection and storage design", 1)
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ("Host", "localhost"),
            ("Port", "5434"),
            ("Database", audit["database"]["name"]),
            ("Database user", audit["database"]["user"]),
            ("PostgreSQL version", audit["database"]["server_version"]),
            ("Current database size", f"{int(audit['database']['size_bytes']) / 1024 / 1024:.2f} MB"),
            ("Password", "Deliberately excluded from this stored document"),
            ("Crop photos", "Private S3 objects; PostgreSQL stores keys and metadata only"),
        ],
        [2400, 6960],
        font_size=9.5,
    )
    add_body(doc, "The local database is an isolated PostgreSQL 16 instance. Native PostgreSQL services on ports 5432 and 5433 are separate and are not used by this project. Production credentials must be different from local development credentials and stored in a secret manager or protected environment configuration.")

    add_heading(doc, "3. Employee access and governance", 1)
    e = audit["employee_summary"]
    add_table(
        doc,
        ["Measure", "Verified value"],
        [
            ("Employee records", e["total"]),
            ("Active / Microsoft-login ready", e["active"]),
            ("Pending company email", e["pending"]),
            ("Personal email present", e["with_personal_email"]),
            ("No email of either type", e["without_any_email"]),
            ("Resolved reporting-manager links", e["resolved_managers"]),
            ("Unresolved manager names", e["unresolved_managers"]),
        ],
        [6100, 3260],
        font_size=9.5,
    )
    admin = audit["admin"]
    add_status_callout(doc, "SOLE ADMINISTRATOR AND APPROVER", f"{admin['full_name']} ({admin['employee_code']}) is active and holds the administrator, manager, product, mapping, expert-review and employee-access approval roles. No other employee has an elevated role.")
    role_rows = [(r["code"], r["name"], r["assigned_employees"], r["description"]) for r in audit["role_counts"]]
    add_table(doc, ["Role code", "Role", "Assigned", "Purpose"], role_rows, [2250, 1800, 900, 4410], font_size=8.5)

    add_heading(doc, "4. Catalogue and mapping status", 1)
    add_body(doc, "All 73 CLSL products are active and approved. Every approval and every catalogue-derived crop/problem mapping is owned by the sole administrator. Mapping rules use only phrases explicitly printed in the supplied CLSL catalogue data; no external product is introduced and no missing dose or registration is invented.")
    category_rows = [(x["name"], x["code"], x["product_count"]) for x in audit["category_counts"]]
    add_table(doc, ["Category", "Code", "Products"], category_rows, [4200, 3500, 1660], font_size=9)

    gaps = [p for p in audit["products"] if p["source_gaps"]]
    add_heading(doc, "4.1 Official-source gaps retained", 2)
    add_body(doc, "The following values are absent from the supplied catalogue dataset. They remain blank intentionally until an authorized CLSL source provides them.")
    add_table(doc, ["Product", "Catalogue fields not supplied"], [(p["name"], p["source_gaps"]) for p in gaps], [3000, 6360], font_size=9)

    add_heading(doc, "5. Database table inventory", 1)
    inventory_rows = [(name, audit["counts"][name], TABLE_PURPOSES.get(name, "System table.")) for name in audit["tables"]]
    add_table(doc, ["Table", "Rows", "Purpose"], inventory_rows, [2600, 900, 5860], font_size=8.5)
    add_body(doc, "Operational tables are expected to contain zero rows until the web application is connected to the FastAPI/PostgreSQL backend. Their schema is ready now so field inspections can be stored server-side rather than only in browser storage.")

    add_heading(doc, "6. Integrity checks", 1)
    labels = {
        "duplicate_employee_codes": "Duplicate employee codes",
        "duplicate_office_emails": "Duplicate office emails",
        "invalid_office_domains": "Office emails outside croplifescience.com",
        "employees_without_field_role": "Employees missing field-employee role",
        "elevated_roles_outside_admin": "Elevated roles assigned outside Jiten",
        "approved_products_without_admin": "Approved products not approved by Jiten",
        "approved_crop_mappings_without_admin": "Approved crop mappings not approved by Jiten",
        "approved_problem_mappings_without_admin": "Approved problem mappings not approved by Jiten",
    }
    add_table(doc, ["Check", "Exceptions", "Result"], [(labels[k], v, "PASS" if int(v) == 0 else "REVIEW") for k, v in integrity.items()], [6500, 1260, 1600], font_size=9)

    add_heading(doc, "7. Complete schema dictionary", 1)
    add_body(doc, "The following sections describe every live public table and column. Required means the database rejects a null value. Defaults shown are database-side defaults; application validation may impose additional rules.")
    by_table = defaultdict(list)
    for column in audit["columns"]:
        by_table[column["table_name"]].append(column)
    constraints = defaultdict(list)
    for item in audit["constraints"]:
        constraints[item["table_name"]].append(item)
    indexes = defaultdict(list)
    for item in audit["indexes"]:
        indexes[item["table_name"]].append(item)

    for table_index, table_name in enumerate(audit["tables"], start=1):
        add_heading(doc, f"7.{table_index} {table_name}", 2)
        add_body(doc, TABLE_PURPOSES.get(table_name, "System table."), bold_lead=f"Rows: {audit['counts'][table_name]}. ")
        field_rows = []
        for column in by_table[table_name]:
            default = column.get("column_default") or "—"
            if len(str(default)) > 42:
                default = str(default)[:39] + "..."
            field_rows.append((column["column_name"], type_label(column), "Yes" if column["is_nullable"] == "NO" else "No", default, describe_column(column)))
        add_table(doc, ["Field", "Type", "Required", "Default", "Meaning"], field_rows, [2050, 1550, 950, 1750, 3060], font_size=7.7)
        cons = constraints[table_name]
        if cons:
            add_heading(doc, "Constraints", 3)
            add_table(doc, ["Type", "Definition"], [(c["constraint_type"], c["definition"]) for c in cons], [1850, 7510], font_size=7.8)
        non_constraint_indexes = [i for i in indexes[table_name] if not i["index_name"].endswith("_pkey")]
        if non_constraint_indexes:
            add_heading(doc, "Indexes", 3)
            add_table(doc, ["Index", "Definition"], [(i["index_name"], i["definition"]) for i in non_constraint_indexes], [3000, 6360], font_size=7.6)

    add_heading(doc, "8. Catalogue taxonomy", 1)
    add_heading(doc, "8.1 Crops and crop groups", 2)
    crop_rows = []
    for item in audit["crop_counts"]:
        aliases = item["aliases"] if isinstance(item["aliases"], list) else []
        crop_rows.append((item["name"], ", ".join(aliases), item["mapped_products"], item["status"]))
    add_table(doc, ["Crop / group", "Catalogue aliases", "Mapped products", "Status"], crop_rows, [2550, 4150, 1400, 1260], font_size=8.3)

    add_heading(doc, "8.2 Problems", 2)
    problem_rows = [(p["issue_type"], p["name"], p["mapped_products"], p["status"]) for p in audit["problem_counts"]]
    add_table(doc, ["Issue type", "Problem", "Mapped products", "Status"], problem_rows, [2200, 4300, 1500, 1360], font_size=8.1)

    add_heading(doc, "9. Full CLSL product inventory", 1)
    add_body(doc, "This inventory confirms the complete 73-product master without reproducing private employee data. Crop/problem counts show catalogue-explicit mappings only.")
    product_rows = [(p["name"], p["category"], p["source_page"], p["crop_mappings"], p["problem_mappings"], p["source_gaps"] or "Complete in supplied dataset") for p in audit["products"]]
    add_table(doc, ["Product", "Category", "Page", "Crop maps", "Problem maps", "Source completeness"], product_rows, [2050, 2250, 650, 900, 1050, 2460], font_size=7.5)

    add_heading(doc, "10. Retention, privacy and training-data rule", 1)
    retention = next((s for s in audit["settings"] if s["key"] == "inspection_image_retention"), None)
    add_status_callout(doc, "IMAGE RETENTION POLICY", "Inspection images remain retained until a replacement model has been trained, validated and made operational. The archive/deletion policy must then be reviewed explicitly. Training consent is stored separately for each image and defaults to false.")
    if retention:
        add_table(doc, ["Setting", "Stored value", "Updated by"], [(retention["key"], json.dumps(retention["value"], ensure_ascii=True), retention["updated_by"])], [2600, 5360, 1400], font_size=8.2)
    add_body(doc, "Employee records contain personal information and must be protected by role-based access, HTTPS, database backups, least-privilege credentials and authorized exports. This document intentionally reports employee data only in aggregate and excludes the database password.")

    add_heading(doc, "11. Backup and recovery reference", 1)
    add_body(doc, "Create encrypted, access-controlled PostgreSQL backups and test restoration before production rollout. Use pg_dump with host localhost, port 5434, database crop_life_ai and user crop_life; enter the password only at the secure prompt. Store production backups outside the application server, with lifecycle and access policies approved by CLSL.")
    add_table(doc, ["Backup component", "Required protection"], [
        ("PostgreSQL", "Encrypted logical/physical backups, retention schedule and restore test"),
        ("S3 crop images", "Private bucket, encryption, versioning/lifecycle and restricted IAM"),
        ("Employee source", "Private import location; never commit to Git"),
        ("Secrets", "Environment/secret manager only; never store in reports or source control"),
        ("Audit evidence", "Retain migration checksums, approval metadata and audit logs"),
    ], [2800, 6560], font_size=9)

    add_heading(doc, "12. Current readiness and next integration", 1)
    add_status_callout(doc, "DATABASE FOUNDATION READY", "The database is ready for backend integration. Product and employee masters are populated; operational tables are present and empty by design until application traffic begins.")
    add_body(doc, "The next implementation step is to connect Microsoft Entra login and the live web inspection flow to FastAPI. Each completed inspection should write the inspection, S3 image metadata, Gemini prediction, recommendation ranking, weather snapshot and AI usage row inside controlled backend workflows.")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 4:
        raise SystemExit("Usage: python build_database_dictionary.py AUDIT.json OUTPUT.docx LOGO.png")
    main(sys.argv[1], sys.argv[2], sys.argv[3])
