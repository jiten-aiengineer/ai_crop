from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Crop_Life_AI_Spray_Calculator_and_Farmer_Cost_Guide.docx"
LOGO = ROOT / "public" / "clsl-logo.png"

GREEN = "1F6B43"
DARK_GREEN = "183326"
LIME = "CBE968"
LIGHT_GREEN = "EAF3D7"
PALE = "F5F7F0"
BLUE = "2E74B5"
MUTED = "66776D"
WHITE = "FFFFFF"
GRID = "C8D4CB"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_fixed_table(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table.rows[0].cells[0]._tc.getparent().getparent()
    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=DARK_GREEN, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_body(doc, text: str, *, bold_lead: str | None = None, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest)
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text: str, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_fixed_table(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        set_cell_shading(header.cells[index], "E8EEF5")
        p = header.cells[index].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=9, bold=True, color=DARK_GREEN)
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_run(p.add_run(value), size=9.2, color=DARK_GREEN)
    set_fixed_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, DARK_GREEN, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header_p.add_run("CROP LIFE AI  |  OPERATIONAL REFERENCE"), size=8, bold=True, color=MUTED)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.columns[0].width = Inches(5.6)
    table.columns[1].width = Inches(0.9)
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    set_run(p.add_run("Crop Life Science Limited  •  Spray Calculator Guide  •  Internal pilot"), size=8, color=MUTED)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Page "), size=8, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(46)
    p.paragraph_format.space_after = Pt(22)
    p.add_run().add_picture(str(LOGO), width=Inches(1.45))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run("OPERATIONAL FIELD GUIDE"), size=9, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("Crop Life AI Spray Calculator"), size=28, bold=True, color=DARK_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run(p.add_run("Quantity, pump mixing, pack planning and farmer cost logic"), size=14, color=BLUE)

    table = doc.add_table(rows=1, cols=3)
    set_fixed_table(table, [3120, 3120, 3120])
    for cell, label, value in zip(
        table.rows[0].cells,
        ("VERSION", "ISSUED", "STATUS"),
        ("1.0", "3 September 2026", "Production pilot"),
    ):
        set_cell_shading(cell, LIGHT_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(label + "\n"), size=8, bold=True, color=GREEN)
        set_run(p.add_run(value), size=10.5, bold=True, color=DARK_GREEN)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Purpose"), size=11, bold=True, color=GREEN)
    add_body(
        doc,
        "This guide records exactly how the web app converts field area, approved dose, water volume, pump size, pack size and price into an application plan and farmer cost estimate.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=18,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Safety boundary: calculation is not a pesticide recommendation. The approved label, crop registration and local expert direction remain authoritative."), size=9.5, bold=True, color="9A4D30")
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. What the calculator does", 1)
    add_body(doc, "The calculator is a planning aid for an already selected and approved CLSL crop-protection product. It does not choose the dose. It scales an approved per-acre dose to the farmer’s actual area, keeps the same product-to-water concentration in every pump, and estimates whole-pack purchase cost.")
    add_table(
        doc,
        ["Input", "Meaning", "Accepted units / rule"],
        [
            ["Area to treat", "Actual field area for this spray", "Acre, hectare or guntha"],
            ["Approved dose", "Label-approved product amount per acre", "ml, g, L or kg per acre"],
            ["Water per acre", "Carrier-water volume", "Litres per acre"],
            ["Pump size", "Usable spray-tank capacity", "Litres"],
            ["Pack size", "One purchasable pack in the same unit as dose", "ml, g, L or kg"],
            ["Pack price", "Farmer’s current price for one selected pack", "Indian rupees"],
            ["Application cost", "Labour and machinery for spraying", "Rupees per acre"],
        ],
        [2100, 4260, 3000],
    )

    add_heading(doc, "2. Area conversion", 1)
    add_body(doc, "All downstream calculations use acres as the common internal area unit.")
    add_table(
        doc,
        ["Selected unit", "Conversion to acres", "Example"],
        [
            ["Acre", "Acres = entered area", "3.5 acre = 3.5 acre"],
            ["Hectare", "Acres = hectares × 2.47105", "1 ha = 2.47105 acre"],
            ["Guntha", "Acres = guntha ÷ 40", "20 guntha = 0.5 acre"],
        ],
        [2100, 3960, 3300],
    )

    add_heading(doc, "3. Product and water quantities", 1)
    add_bullet(doc, "Total product = area in acres × approved dose per acre")
    add_bullet(doc, "Total water = area in acres × water litres per acre")
    add_body(doc, "The product unit is preserved. A dose entered in ml/ac produces a total in ml; g/ac produces g; L/ac produces L; and kg/ac produces kg.")

    add_heading(doc, "4. Pump-load and mixing logic", 1)
    add_body(doc, "The app keeps the product concentration consistent across full and partial tanks. This is safer and more accurate than dividing the product equally by a rounded number of refills.")
    add_number(doc, "Pump loads = round up(total water ÷ pump size).")
    add_number(doc, "Product concentration = total product ÷ total water.")
    add_number(doc, "Product in a full tank = concentration × pump size.")
    add_number(doc, "Final tank water = total water − pump size × (pump loads − 1).")
    add_number(doc, "Final tank product = concentration × final tank water.")
    add_body(doc, "If total water divides exactly by pump size, the final tank is a normal full tank. If not, the last tank shows the smaller remaining water and matching product amount.", bold_lead="If total water divides exactly by pump size")

    add_heading(doc, "5. Pack purchase and farmer cost", 1)
    add_bullet(doc, "Packs to buy = round up(total product ÷ pack size).")
    add_bullet(doc, "Pack purchase cost = packs to buy × price per pack.")
    add_bullet(doc, "Application cost = area in acres × labour/machinery cost per acre.")
    add_bullet(doc, "Total spray cost = pack purchase cost + application cost.")
    add_body(doc, "The displayed purchase cost is cash required to buy whole packs. It can be higher than the value of product consumed because the remaining product stays available for a later approved use.")

    add_heading(doc, "6. Worked spray example", 1)
    area = 3.5
    dose = 80
    water_per_acre = 125
    tank = 15
    pack_size = 200
    pack_price = 550
    application_per_acre = 250
    total_product = area * dose
    total_water = area * water_per_acre
    loads = math.ceil(total_water / tank)
    concentration = total_product / total_water
    full_tank_product = concentration * tank
    final_water = total_water - tank * (loads - 1)
    final_product = concentration * final_water
    packs = math.ceil(total_product / pack_size)
    purchase_cost = packs * pack_price
    application_cost = area * application_per_acre
    total_cost = purchase_cost + application_cost
    add_table(
        doc,
        ["Step", "Calculation", "Result"],
        [
            ["Product required", f"{area} acre × {dose} ml/ac", f"{total_product:.0f} ml"],
            ["Water required", f"{area} acre × {water_per_acre} L/ac", f"{total_water:.1f} L"],
            ["Pump loads", f"Round up({total_water:.1f} ÷ {tank})", str(loads)],
            ["Full tank mix", f"({total_product:.0f} ÷ {total_water:.1f}) × {tank}", f"{tank} L + {full_tank_product:.1f} ml"],
            ["Final tank mix", f"Remaining water × concentration", f"{final_water:.1f} L + {final_product:.1f} ml"],
            ["Packs to buy", f"Round up({total_product:.0f} ÷ {pack_size})", f"{packs} packs"],
            ["Pack purchase", f"{packs} × ₹{pack_price}", f"₹{purchase_cost:,}"],
            ["Application cost", f"{area} × ₹{application_per_acre}", f"₹{application_cost:,.0f}"],
            ["Total spray cost", f"₹{purchase_cost:,} + ₹{application_cost:,.0f}", f"₹{total_cost:,.0f}"],
        ],
        [2150, 4510, 2700],
    )

    add_heading(doc, "7. Catalogue prefill behavior", 1)
    add_body(doc, "When the catalogue contains one clear numeric dose tied directly to acre or hectare, the app may prefill it. Hectare doses are converted to acres by dividing by 2.47105. Range doses and crop-dependent instructions are not silently averaged. The interface marks a prefilled value and asks the user to verify it.")
    add_bullet(doc, "The printed catalogue text remains visible next to the input.")
    add_bullet(doc, "Changing the dose manually removes the prefill indicator.")
    add_bullet(doc, "Pack size may be seeded from the first compatible catalogue packing size, but the farmer must enter the current price.")
    add_bullet(doc, "No price is inferred by AI and no product dose is invented.")

    add_heading(doc, "8. Farm cost planner", 1)
    add_body(doc, "The separate farm cost planner estimates the combined field-operation cost. The fertilizer calculator has been removed because Crop Life Science does not offer a general fertilizer service in this tool.")
    add_bullet(doc, "Per-acre field cost = seed/planting + irrigation + crop protection + labour/machinery + harvest/transport.")
    add_bullet(doc, "Estimated total field cost = per-acre field cost × area in acres.")
    add_body(doc, "These are farmer-entered estimates, not quotations. Market price, labour availability, crop stage and local practice can change the actual cost.")

    add_heading(doc, "9. Operational checks before use", 1)
    for item in (
        "Confirm the exact crop and target problem on the approved product label.",
        "Confirm the dose unit: ml, g, L or kg per acre.",
        "Use the field’s calibrated water volume rather than a generic default where available.",
        "Confirm the pump’s usable capacity and recalibrate worn nozzles.",
        "Check live weather guidance; avoid rain risk, strong wind and unsafe heat.",
        "Wear label-required personal protective equipment and follow re-entry guidance.",
        "Treat the calculated amount as a plan; a qualified local expert and approved label remain authoritative.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "10. Validation test cases", 1)
    add_table(
        doc,
        ["Case", "Expected check", "Pass condition"],
        [
            ["Zero dose", "No product should be implied", "Product required = 0 and warning visible"],
            ["Exact tank division", "No under/over-mixing", "Final tank equals one full tank"],
            ["Partial final tank", "Concentration remains constant", "Final product scales to remaining water"],
            ["Small pack", "Whole-pack cash cost", "Packs are rounded upward"],
            ["Hectare / guntha", "Common internal area", "Converted acres drive all totals"],
            ["Catalogue range dose", "No silent midpoint", "User must confirm approved dose"],
        ],
        [2100, 4020, 3240],
    )

    add_heading(doc, "11. Governance and change control", 1)
    add_body(doc, "Any formula, unit conversion, catalogue parsing or cost-field change should update both the application and this guide in the same release. Product label data remains governed by the CLSL catalogue and the approval workflow in the product database.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell = doc.add_table(rows=1, cols=1)
    set_fixed_table(set_cell, [9360])
    cell = set_cell.cell(0, 0)
    set_cell_shading(cell, DARK_GREEN)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Crop Life AI  •  Practical crop protection support, grounded in the CLSL catalogue"), size=10, bold=True, color=WHITE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
